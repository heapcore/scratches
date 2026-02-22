import docx
import StringIO


class DocxHandler(object):
    def __init__(self, input_file):
        self.input_file = input_file
        self.document = docx.Document(self.input_file)

    def save(self, output_file=None, return_content=False):
        if not output_file and not return_content:
            raise ValueError(
                "You must pass output_file or return_content=True to the function"
            )

        output_content = StringIO.StringIO()
        if not return_content:
            self.document.save(output_file)
        else:
            self.document.save(output_content)

        if return_content:
            return output_content

    def format(self, **kwargs):
        def replace_in_paragraph(paragraph, pattern, replace_text):
            if not isinstance(replace_text, basestring):
                replace_text = str(replace_text)
            runs = paragraph.runs
            for i in range(len(runs)):
                if pattern in runs[i].text:
                    text = runs[i].text.replace(pattern, replace_text)
                    runs[i].text = text

        for paragraph in self.document.paragraphs:
            for key in kwargs:
                pattern = "{" + key + "}"
                if pattern in paragraph.text:
                    replace_in_paragraph(paragraph, pattern, kwargs[key])

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key in kwargs:
                            pattern = "{" + key + "}"
                            if pattern in paragraph.text:
                                replace_in_paragraph(paragraph, pattern, kwargs[key])


doc = DocxHandler("test.docx")
doc.format(day=17, month=12, year=1900, str_data="test")
doc.save("test_output.docx")
